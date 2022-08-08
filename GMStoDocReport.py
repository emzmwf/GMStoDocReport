# to be run within GMS
import DigitalMicrograph as DM

'''from PIL import Image, ImageEnhance'''
# Note, these were installed in an earlier version to apply python filters
# and left in in case I ever want to use them again
# 

import numpy as np

##Path handling via pathlib
from pathlib import Path

# To add path names
import os

# use docx to merge and produce word document

import docx
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docxcompose.composer import Composer

####
DM.ClearResults()
print("\n======DM To Report ======\n")

##Check main thread
if (DM.IsScriptOnMainThread() == False):
	print( ' MatplotLib and scipy scripts require to be run on the main thread.',
		'\n Uncheck the "Execute on Background Thread"',
		'checkbox at the bottom of the Script Window' )
	exit()

	
#########################################################
## v 1.4    MWF August 2022
##
##  Location of templates stored in a global library
##  Ask the user where they are if library is not found
##
########################################################

#########################################################
##  To do list v1.5
##  
##  Check required libraries are installed and suggest installation method
##
##  Check the templates we need are where the library says they are
##  Dialog box before asking for template folder
##  
##  TIF output version
##  Create temporary image without needing to show it and save it?
################################################################

##########################################################
#DEF for selecting a directory
#
#This appears overly inconvenient
#But it works

def WhatFolder():
    dmScriptF = 'string folder , outputFolder' + '\n'
    dmScriptF += 'if ( !GetDirectoryDialog( "Select ouput folder" , "" , folder ) ) ' + '\n'
    dmScriptF += '     Result(folder)' + '\n'
    dmScriptF += '     string Dir = folder' + '\n'
    dmScriptF += '     string Dir2 = folder' + '\n'
    dmScriptF += '     TagGroup tg = GetPersistentTagGroup( ) ' + '\n'
    dmScriptF += '     tg.TagGroupSetTagAsString( "DM2Python String", folder )' + '\n'

    #Execute the script
    DM.ExecuteScriptString( dmScriptF )

    #Get the selection data into python
    TGp = DM.GetPersistentTagGroup()
    returnVal, val = TGp.GetTagAsText('DM2Python String')
    #val is the python string containing the folder
    return(val)

##End of what folder def




##Dialog to ask if one file or workspace
bVal = DM.OkCancelDialog( 'Select OK for all workspace or cancel for front image only' )
if (bVal == True):
    choice = "Workspace"
else:
    choice = "Front Image"

print( 'You have selected: ', choice )

###Location of templates
#Get the persistant tag group
TGp = DM.GetPersistentTagGroup()
#See if the tag we are after exists
returnVal, val = TGp.GetTagAsText('DM2ReportTemplateFolder')


#True if it is there
if returnVal:
    print("found the templates")
    Spath = val
    
#False if it isn't there
if not returnVal:    
    print("where are the templates?")
    Spath = WhatFolder()
    print("you say there are here: ")
    print(Spath)
    
    bVal = DM.OkCancelDialog( 'Save this location to global tags for next time?' )

    if (bVal == True):
        choice = "Save"
        TGp.SetTagAsText('DM2ReportTemplateFolder',Spath)
    else:
        choice = "Don't save"

##Now we have the template location


#Ask for where the report is to go
DM.OkDialog('Select the folder to save the report in') 
Outpath = WhatFolder()
print("Report will be saved in the folder ")
print(Outpath)

##name for workspace data


###############################################
## More Definitions
#################################################
##Def - input is a DM tag group, output is a dictionary
def ImportTags(tagGroup):
    #List of tags, must be in same order as StrList
    TagList = ["AnDate",
    "Instrument",
    "Detector",
    "mag",
    "StageX",
    "StageY",
    "StageZ"
    ]

    #List of tags, must be in same order as TagList
    StrList = ["DataBar:Acquisition Date",
    "Session Info:Microscope",
    "DataBar:Signal Name",
    "Microscope Info:Formatted Indicated Mag",
    "Microscope Info:Stage Position:Stage X",
    "Microscope Info:Stage Position:Stage Y",
    "Microscope Info:Stage Position:Stage Z"]
    

    ##Dictionary of the pairs, note these don't have an order, order
    ##is set here by the TagList and StrList
    lstdict = {
    "AnDate": 20000101,
    "Instrument": "TestTEM",
    "Detector": "Virtual",
    "mag": 1,
    "StageX": 0,
    "StageY": 0,
    "StageZ": 0
    }
    LooLen = len(lstdict)
    for x in range(LooLen):
        Tstring = StrList[x]
        returnVal, val1 = tagGroup.GetTagAsString(Tstring) 
        dicval = TagList[x]
        lstdict[dicval] = val1
    
    return (lstdict)
##End of def ImportTags


##
def DoIt():
        global dmImg
        #global $path
        #dmImg = DM.GetFrontImage() 
        type = dmImg.GetDataType() 
        if (type != 23):
            dmImgData = dmImg.GetNumArray() # Get NumpyArray to image data

        #Get type of image
        #Make a copy that is int8 if it isn't already
        #We need int8 to create figure
        #And DM does this better than PIL

        #print(type(dmImg))
        #print("dmImgData is type:")
        #print(type(dmImgData))  #should be a numpy.ndarray 
        
        ####################################
        ##Get metadata from the image
        ###################################


        ##temp data for testing
        figname="TEST"


        figname = dmImg.GetName()
        tagGroup = dmImg.GetTagGroup() 

        #Turn tag group into dictionary using a def ImportTags(tag group)
        lstdict = ImportTags(tagGroup)

        AnDate = (lstdict["AnDate"])
        Instrument =(lstdict["Instrument"])
        Detector = (lstdict["Detector"])
        mag = (lstdict["mag"])
        StageX = (lstdict["StageX"])
        StageY = (lstdict["StageY"])

	'''
        ##################################################################
        ## Commented out section - 					##
	## This was to allow any filtering for a template to be run here##
        ##################################################################

        ###Create image from DM data via PIL
        if (type !=23):
            image1 = Image.fromarray(dmImgData)
            ImMod = (image1.mode)   #Get this info
            # F is 32 bit floating point, tif save only
            # RGB is 3x8-bit pixels, true color

            #print(image1.mode)  #F means float
            #print(image1.size)
            wid = image1.size[0]
            hgt = image1.size[1]

            ##maximum value of array)
            maxElement = np.amax(dmImgData)
            #print(maxElement)


            dmImg2 = DM.CreateImage(dmImgData)
        '''
        if (type == 23):
            print("RGB image found")
            ##get width and height
            wid = dmImg.GetImgWidth()
            hgt = dmImg.GetImgHeight()
            

        ##Change image to an rgb version with the scale bar still seen
        dmScript = '//This is a DM script' + '\n'
        dmScript+= 'image img3 := GetFrontImage()'  + '\n'
        dmScript+= 'rgbimage RGBimg =createimagefromdisplay(img3)'  + '\n'
        dmScript+= 'SetName(RGBimg, "RGB Image")'  + '\n'
        dmScript+= 'ShowImage (RGBimg)'

        DM.ExecuteScriptString(dmScript)
        ###now get this new image
        dmImg2 = DM.GetFrontImage() # Get reference to front most image
        dmImg2.UpdateImage()

	wid = dmImg2.GetImgWidth()
        hgt = dmImg2.GetImgHeight()
	
	
        ##Get today's date            
        from datetime import datetime
        Today = datetime.today().strftime('%Y%m%d')


        ##Define the end filename
        #OutName = 'nmRC_TEM_Report_'+figname+'.docx'
        Hfile = Path(Spath,'Header.docx')
        ##$path is defined outside of this def
        ##get bVal and set OutFile based on one file or workspace
        
        if (bVal ==True):
            ##workspace
            
            global OutName
            
        
            
            ##check Wfile, first or not first
            if (Wfile == "First"):
                OutName = 'GMS_TEM_Report.docx'
                OutFile = Path(Outpath,OutName)     #v1.3 - where the user has selected
                WOutFile = OutFile
                Hfile = Path(Spath,'Header.docx')
                ##Check if the file already exists
                A=1
                while (OutFile.exists()):
                    OutName = 'GMS_TEM_Report_'+str(A)+'.docx'
                    OutFile = Path(Outpath,OutName)       #v 1.3 onwards - folder selected by dialog
                    WOutFile = OutFile
                    A = A+1       
                              
            else:
                print(OutName)
                OutFile = Path(Outpath, OutName)    #Save to the user selected folder
                Hfile = OutFile
            print(OutFile)     
        else:    
            ##Single file
            OutName = 'GMS_TEM_Report_'+figname+'.docx'
            OutFile = Path(Outpath,OutName)      #Save to the user selected folder
            Hfile = Path(Spath,'Header.docx')
            #############################
            ##For single frame output
            ##Check if the outfile already exists / is open
            B = 1 ##I find the letter B visible and friendly
            while (OutFile.exists()):
                OutName = 'GMS_TEM_Report_'+figname+'_'+str(B)+'.docx'
                OutFile = Path(Outpath,OutName)   #v1.3 on, where user has said
                B = B+1


        ##Take style from doc1
        tdocument = Document(Hfile)
        Ostyles = tdocument.styles
        section = tdocument.sections[0]
        header = section.header

        ##Create new document
        ##use Header.docx for single file, or first file in workspace
        ##use the workspace file otherwise
        ##Now defining Hfile earlier   Hfile = Path(Spath,'Header.docx')
        mdoc = docx.Document(Hfile)

        paragraph = mdoc.add_paragraph(figname)

        ##do we have to save the image? It appears we do

        ##change the directory
        import os
        Tfile2 = os.path.join(Spath, 'Tfile2.tif')
        try:
            #dmImg2.SaveImage(Tfile2)
            ##################
            ##How to fix######
            ##################
            #Get the image document that dmImg2 is in
            # Save the image document, e.g 
            #ImageDocA.SaveToFile('TIFF Format', 'C:/Users/emzmwf/Temp/ABC.tif')
            #################################################################
            
            ImageDocA = DM.GetFrontImageDocument()
            ImageDocA.SaveToFile('TIFF Format', Tfile2)
         
            
        except IOError:
            print("Could not open file, still in use?")
            
            
        DM.DeleteImage(dmImg2)   ##removes it from memory without asking
        DM.CloseImage(dmImg2)   ##closes the file so it doesn't trigger a cannot access error - run after Delete to avoid
        ##getting a dialog box

        ##width and height in inches for printing
        winch = 5
        hinch = (winch/wid)*hgt

        mdoc.add_picture(Tfile2, width = Inches(winch), height = Inches(hinch))

        #############################
        ##Replace keywords in parts

        ##meta.docx contains the meta details
        ##Endtext.docx is end of document information

        files = ['meta.docx', 'Endtext.docx']

        def combine_word_documents(files):
            composer = Composer(mdoc)    ##create a composer for this document
            ##Define some styles
            styles = mdoc.styles
            styles['Normal'].font.name = 'Calibri'
            styles['Normal'].font.size = Pt(10)
                    
            for index, file in enumerate(files):
                Tfile3 = os.path.join(Spath, file)  ##Note file, NOT files
                tdocument = Document(Tfile3)
                sub_doc = Document(Tfile3)
                
                ##Replace keywords
                for table in sub_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if '$AnDate' in paragraph.text:
                                    paragraph.text = paragraph.text.replace("$AnDate", AnDate)
                                    #print("Replaced Andate")
                                if '$TEM' in paragraph.text:
                                    paragraph.text = paragraph.text.replace("$TEM", Instrument)
                                if '$DET' in paragraph.text:
                                    paragraph.text = paragraph.text.replace("$DET", Detector)
                                if '$StX' in paragraph.text:
                                    paragraph.text = paragraph.text.replace("$StX", StageX)
                                if '$StY' in paragraph.text:
                                    paragraph.text = paragraph.text.replace("$StY", StageY)                            
                                if '$MAG' in paragraph.text:
                                    paragraph.text = paragraph.text.replace("$MAG", mag)                                     
                
                ##Append the updated doc
                composer.append(sub_doc)
            ##Add page break if we need it
            if (Wfile == "NOT FIRST"):
                mdoc.add_page_break()
            ##Save the doc
            mdoc.save(OutFile)
        ##End of def combine_word_documents

        combine_word_documents(files)
        #print("Created new file: "+OutName)
        
        del dmImg   #remove python thing from memory
        del dmImg2      #remove python thing from memory
        del Tfile2
        return

# run front image or all in workspace
#bVal = False, single image


###############################################
## End of definitions, now run the rest of the program
#################################################

if (bVal == False):
    dmImg = DM.GetFrontImage() 
    Wfile = ""
    DoIt()
    print("\n File name is: \n")
    print(OutName)
        
else:
    #Get current Document, and then the workspace it is on
    ImgD = DM.GetFrontImageDocument() 
    dmImgW = ImgD.GetImage(0)
    wsID = ImgD.GetWorkspace() 
    
    ##How many images on this workspace
    Ino = DM.CountImageDocuments(wsID)
    ##Create array to hold the list of image IDs
    ImArr = []
    Wfile = ""
    
    ##Loop through the number of images
    for i in range(0, Ino):
        nom = ""
        if (i==0):
            dmImgW = DM.GetFrontImage()
            nom = dmImgW.GetName()
            imageID = dmImgW.GetID()  
            ImArr.append(imageID)
            #print(nom)
        else:
            dmImgW = dmImgW.FindNextImage()
            nom = dmImgW.GetName()
            imageID = dmImgW.GetID()
            ImArr.append(imageID)

        imgGetName = dmImgW.GetName()
        imageID = dmImgW.GetID()  
        dmImg = dmImgW
        
    ##Now loop through the image list, show each image, run DoIt    
    Fval = ImArr[0]
    OutName = ""
    for x in ImArr:
        dmImg = DM.FindImageByID(x)
        dmImg.ShowImage()
        if (x == Fval):
            Wfile = "First"
        else:
            Wfile = "NOT FIRST"
        DoIt()
    print("\n File name is: \n")
    print(OutName)
    del dmImgW
##end of dir


